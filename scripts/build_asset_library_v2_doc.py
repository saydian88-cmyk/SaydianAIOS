from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"F:\xcodeplace\工作督查\Saydian AI 素材库 V2.0 落地设计与实施方案.docx")
SOURCE = Path(r"F:\xcodeplace\工作督查\Saydian AI 素材库 V1.0 实施规范.docx")
REPO = Path(r"F:\xcodeplace\工作督查\saidian-ops-center")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
RED = "9B1C1C"
GREEN = "1F5C42"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_fonts(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def configure_table(table, widths):
    assert sum(widths) == TABLE_WIDTH
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    configure_table(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_fonts(run, 9.5, True, NAVY)
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.15
            run = paragraph.add_run(str(value))
            set_fonts(run, 9.2, False, "222B3A")
        configure_table(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_numbering(doc):
    numbering = doc.part.numbering_part.element

    def create_num(abstract_id, num_id, fmt, text_value):
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        level.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        lvl_text.set(qn("w:val"), text_value)
        level.append(lvl_text)
        suff = OxmlElement("w:suff")
        suff.set(qn("w:val"), "tab")
        level.append(suff)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "540")
        ind.set(qn("w:hanging"), "271")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        level.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Calibri")
        fonts.set(qn("w:hAnsi"), "Calibri")
        r_pr.append(fonts)
        level.append(r_pr)
        abstract.append(level)
        numbering.append(abstract)
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)

    create_num(91, 91, "bullet", "•")
    create_num(92, 92, "decimal", "%1.")


def list_item(doc, text_value, ordered=False):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(text_value)
    set_fonts(run, 11, False, "222B3A")
    return paragraph


def add_heading(doc, text_value, level):
    paragraph = doc.add_heading(text_value, level=level)
    return paragraph


def add_body(doc, text_value, bold_prefix=None):
    paragraph = doc.add_paragraph()
    if bold_prefix and text_value.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_fonts(first, 11, True, NAVY)
        rest = paragraph.add_run(text_value[len(bold_prefix):])
        set_fonts(rest, 11, False, "222B3A")
    else:
        run = paragraph.add_run(text_value)
        set_fonts(run, 11, False, "222B3A")
    return paragraph


def add_callout(doc, title, body, fill=LIGHT_GRAY, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    configure_table(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title)
    set_fonts(run, 11, True, color)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    run2 = p2.add_run(body)
    set_fonts(run2, 10.5, False, "263043")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_page_break(doc):
    return None


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def configure_header_footer(doc):
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    left = p.add_run("SAYDIAN · 品牌数据中心")
    set_fonts(left, 8.5, True, MUTED)
    right = p.add_run("    素材库 V2.0")
    set_fonts(right, 8.5, False, MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_after = Pt(0)
    label = fp.add_run("内部实施文件  |  ")
    set_fonts(label, 8.5, False, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    fp._p.append(field)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    kicker = p.add_run("SAYDIAN AI OS · BRAND DATA CENTER")
    set_fonts(kicker, 10, True, BLUE)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("Saydian AI 素材库 V2.0")
    set_fonts(run, 25, True, NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("落地设计与实施方案")
    set_fonts(run, 16, False, DARK_BLUE)

    add_table(doc, ["项目", "内容"], [
        ("版本", "V2.0"),
        ("日期", "2026-07-23"),
        ("适用系统", "赛电品牌数据中心：品牌知识库 + 素材库"),
        ("技术栈", "NestJS、Prisma、PostgreSQL、Vue 3、PostgreSQL任务队列"),
        ("对象存储", "阿里云 OSS 深圳地域 · saidian-brand-assets-prod-sz"),
        ("代码仓库", "saydian88-cmyk/SaydianAIOS"),
        ("原规范", str(SOURCE)),
    ], [2200, 7160])

    add_callout(doc, "本期范围", "只改造品牌数据中心的品牌知识库与素材库。店铺、发布、直播、客服、竞品等模块保持不变；只预留素材使用和效果回流接口。", "F4F6F9", NAVY)
    add_callout(doc, "核心升级", "从“文件夹 + 表单字段”升级为“素材对象图谱 + 受控标签 + AI处理流水线”。OSS只保存文件，业务分类全部进入数据库。", "E8EEF5", DARK_BLUE)


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_header_footer(doc)
    add_title_page(doc)

    doc.add_page_break()
    add_heading(doc, "一、执行摘要", 1)
    add_body(doc, "V1 的分类、审核、版权、视频模块和使用台账继续保留。V2 不覆盖原规范，而是在现有品牌数据中心代码上实施结构化升级。")
    for item in [
        "一个文件与一项素材分离：逻辑素材可拥有多个物理版本、预览文件、分析产物和派生片段。",
        "原始文件只读且永久保留原始文件名；公开编号不含型号，型号通过多对多关系管理。",
        "处理状态、审核状态、可用状态和权限状态分开，避免“上传成功”被误认为“AI已完成”或“可商用”。",
        "精确重复自动复用主素材并保留上传行为；近似重复仅生成候选，不自动删除或合并。",
        "AI能力以实际配置和执行结果为准。百炼未配置时明确显示“未配置”，本地参数提取不冒充AI识别。",
        "员工上传、AI派生、审核、调用、发布效果和失败原因均进入每日台账。",
    ]:
        list_item(doc, item)

    add_heading(doc, "1.1 成功标准", 2)
    add_table(doc, ["目标", "验收结果"], [
        ("可追溯", "每项素材可追溯上传人、批次、版本、OSS对象、审核与操作记录。"),
        ("可检索", "按型号、类型、层级、模块、人群、场景、功能、平台、市场、员工、权限、评分和状态组合检索。"),
        ("可调用", "只有 APPROVED + ACTIVE + COMMERCIAL/EDIT_ONLY 的素材进入自动内容调用池。"),
        ("可复盘", "员工新增、重复、AI派生、审核通过、实际调用与效果数据在日报中分开统计。"),
        ("可演进", "不移动现有OSS对象，不覆盖V1文档，旧接口兼容一个版本周期。"),
    ], [1800, 7560])

    add_heading(doc, "1.2 明确不做", 2)
    for item in ["不重建商城、客服帮助网站或ERP。", "不改内容生产、平台发布、直播、客服、店铺和竞品模块。", "不引入 Redis、Dify、n8n、独立向量数据库。", "不把已登录账号、已写适配器或本地参数识别显示为真实AI能力已打通。"]:
        list_item(doc, item)

    add_page_break(doc)
    add_heading(doc, "二、总体架构", 1)
    add_table(doc, ["输入", "品牌数据中心 V2", "输出"], [
        ("Web/手机批量上传\n本地只读扫描\n历史OSS对象", "接收与去重 → 版本/关系 → AI任务 → 人工审核 → ACTIVE素材池", "素材检索\n视频模块\n员工台账\n缺口与日报"),
        ("品牌资料\n产品与SKU\n证据和表述规则\n客服知识与FAQ", "品牌版本 → 产品关系 → 知识审核 → FAQ标准化 → AI调用资格", "可信知识检索\n短回复/详细回复\n证据引用"),
    ], [2400, 4560, 2400])
    add_body(doc, "主数据和任务状态保存于一套 PostgreSQL。深圳 OSS 是唯一受管文件库；系统部署包进入 system-deploy/，不进入素材统计。")

    add_heading(doc, "2.1 关键边界", 2)
    add_table(doc, ["对象", "负责内容", "不负责内容"], [
        ("OSS", "原件、派生文件、预览、分析产物、部署包", "业务分类、审核、型号、员工绩效"),
        ("PostgreSQL", "素材对象、版本、关系、标签、任务、审核、台账、效果", "保存大文件本体"),
        ("品牌数据中心", "统一入口、检索、审核、AI任务和日报", "替代商城、ERP、客服站"),
        ("AI适配器", "提出标签、OCR、转写、场景与模块建议", "未经审核修改品牌事实或直接授权商用"),
    ], [1600, 3880, 3880])

    add_page_break(doc)
    add_heading(doc, "三、基础知识体系", 1)
    add_heading(doc, "3.1 品牌版本", 2)
    add_body(doc, "BrandProfileVersion 保存品牌定位、故事、语气、视觉、国家版本、审核人和生效时间。AI只能读取已生效版本；修改生成新版本，不覆盖历史。")
    add_heading(doc, "3.2 产品与SKU", 2)
    add_body(doc, "Product 与 ProductSku 继续作为型号和SKU主数据。素材通过 AssetProduct 关联一个或多个产品；上传和编辑时必须从产品库选择，禁止自由文本产生新型号。")
    add_heading(doc, "3.3 知识与FAQ", 2)
    add_table(doc, ["对象", "关键字段", "AI调用条件"], [
        ("KnowledgeEntry", "分类、标题、正文、来源等级、有效期、证据编号、审核状态", "READY + externallyUsable + 未过期"),
        ("FaqEntry", "标准问题、短回复、详细回复、意图、市场、频次、适用产品", "已审核、对外可用、未过期"),
        ("FaqVariant", "不同问法、来源、出现频次", "继承FAQ主记录"),
        ("EvidenceClaim", "证据编号、覆盖对象、允许表述、限制、有效期", "READY 且有效期内"),
        ("ProductMapping", "商品名、铭牌型号、注册型号、注册编号、发布前动作", "人工确认后引用"),
        ("PhraseRule", "拦截表述、替代表述、使用条件", "active=true"),
    ], [1700, 4800, 2860])

    add_heading(doc, "3.4 来源等级", 2)
    add_table(doc, ["等级", "来源", "默认用途"], [
        ("A", "官方说明、有效证书、检测报告", "产品参数和强事实依据"),
        ("B", "公司内部确认资料", "品牌、产品、服务标准"),
        ("C", "权威媒体或专业机构", "知识解释和背景"),
        ("D", "用户反馈、客服与平台评论", "需求、痛点、FAQ信号"),
        ("E", "竞品宣传或网络参考", "内部研究，不直接作为产品事实"),
    ], [1100, 4700, 3560])

    add_page_break(doc)
    add_heading(doc, "四、素材对象图谱", 1)
    add_heading(doc, "4.1 核心模型", 2)
    add_table(doc, ["模型", "职责"], [
        ("Asset", "逻辑素材。保存公开编号、类型、层级、名称、四类状态、质量和业务说明。"),
        ("AssetVersion", "物理文件版本。保存OSS对象、哈希、尺寸、时长、编码、缩略图和技术元数据。"),
        ("AssetProduct", "素材与产品/型号的多对多关系，保存范围、置信度和人工确认。"),
        ("TagDefinition / AssetTag", "受控标签及其来源、置信度、人工确认、锁定状态和模型版本。"),
        ("AssetRelation", "原片、片段、派生、预览、精确重复和近似重复关系。"),
        ("UploadBatch / UploadEvent", "批量上传、员工归属、精确重复、失败和贡献统计。"),
        ("AssetAnalysisJob", "技术检测、缩略图、OCR、转写、关键帧、切段、理解、标签和近似去重任务。"),
        ("AssetSegment", "视频时间范围、转写、模块类型、分析版本、锁定与物化状态。"),
        ("AssetReviewDecision", "通过、退回、仅内部、拒绝的人工决策记录。"),
        ("AssetUsage / AssetMetricSnapshot", "脚本、成片、AI任务的调用记录及播放、互动、咨询、订单和成交回流。"),
        ("AssetGapSnapshot", "按日期、型号、素材类型和覆盖项保存缺口与补拍建议。"),
    ], [2600, 6760])

    add_heading(doc, "4.2 状态分离", 2)
    add_table(doc, ["状态维度", "枚举", "回答的问题"], [
        ("处理状态", "RECEIVED / HASHED / STORED / ANALYZING / READY_FOR_REVIEW / FAILED", "文件和自动任务处理到哪一步？"),
        ("审核状态", "PENDING / APPROVED / RETURNED / REJECTED", "人工是否批准？"),
        ("可用状态", "INACTIVE / ACTIVE / SUSPENDED / ARCHIVED", "自动内容任务现在能否调用？"),
        ("权限状态", "COMMERCIAL / INTERNAL / EDIT_ONLY / AUTH_REQUIRED / EXPIRED / PROHIBITED", "以什么条件使用？"),
    ], [1700, 5000, 2660])
    add_callout(doc, "调用硬规则", "只有 reviewStatus=APPROVED、availabilityStatus=ACTIVE，且 rightsStatus 为 COMMERCIAL 或 EDIT_ONLY 的素材，才进入自动内容调用池。", "FFF4E5", RED)

    add_heading(doc, "4.3 编号和文件名", 2)
    add_body(doc, "公开编号：SD-{素材类型}-{YYYYMMDD}-{6位随机码}。型号不进入编号；素材可关联多个型号。原始文件名永久保留，只有导出时才生成规范展示文件名。")
    add_table(doc, ["示例", "说明"], [
        ("SD-IMAGE-20260723-A1B2C3", "图片逻辑素材公开编号"),
        ("SD-VIDEO-20260723-D4E5F6", "视频逻辑素材公开编号"),
        ("W9_HOOK_父母关怀_03S_001_FINAL.mp4", "人工导出展示名，不改变OSS原件对象键"),
    ], [3600, 5760])

    add_page_break(doc)
    add_heading(doc, "五、受控标签体系", 1)
    add_body(doc, "标签由 TagDefinition 管理，不能让不同员工随意产生同义词。AI标签保留模型版本和置信度；人工修改后 locked=true，后续分析不得覆盖。")
    add_table(doc, ["命名空间", "示例", "规则"], [
        ("asset_category", "白底图、主图、详情图、场景图、佩戴图、功能图", "图片/文件用途分类"),
        ("video_module", "HOOK、PAIN、FEATURE、DEMO、CTA", "使用固定枚举，不使用自由文本"),
        ("scene", "家庭、运动、通勤、送礼、直播", "允许逐步扩充，需管理员启用"),
        ("audience", "中老年、成年子女、女性、运动人群", "描述目标人群，不与人物识别混用"),
        ("feature", "血压、心率、血氧、睡眠、GPS、SOS", "需与产品和证据关系核对"),
        ("pain_point", "父母关怀、睡眠问题、测量不便、续航顾虑", "来自内容和用户问题"),
        ("platform", "抖音、视频号、小红书、TikTok", "表示适配平台，不表示已经发布"),
        ("market", "CN、US、GLOBAL", "市场范围"),
        ("language", "zh-CN、en-US", "内容语言"),
        ("source", "员工拍摄、供应商、UGC、AI、竞品参考", "必须与权限状态同时保存"),
    ], [1900, 3500, 3960])

    add_heading(doc, "5.1 视频模块枚举", 2)
    add_table(doc, ["模块", "作用", "典型位置"], [
        ("HOOK", "前1至3秒吸引注意", "开头"), ("PAIN", "表达用户痛点", "开头/中段"),
        ("SCENE", "建立使用场景", "开头/中段"), ("FEATURE", "展示产品功能", "中段"),
        ("BENEFIT", "把功能翻译为用户价值", "中段"), ("PROOF", "提供资料、评价或证据", "中段"),
        ("DEMO", "操作和使用演示", "中段"), ("COMPARE", "新旧或方案对比", "中段"),
        ("UGC", "用户体验与评价", "中段"), ("STORY", "人物或情节叙事", "全段"),
        ("TRANSITION", "镜头衔接", "中段"), ("TRAFFIC", "评论、关注、私信等引流", "结尾"),
        ("OFFER", "优惠和活动", "结尾"), ("CTA", "购买或行动引导", "结尾"), ("ENDING", "品牌收尾", "结尾"),
    ], [1500, 5360, 2500])

    add_page_break(doc)
    add_heading(doc, "六、深圳 OSS 存储设计", 1)
    add_table(doc, ["目录", "用途"], [
        ("brand-assets/original/{hash前2位}/{sha256}.{ext}", "原始文件。按内容寻址，禁止业务重命名。"),
        ("brand-assets/derived/{assetId}/{派生类型}/{version}/{sha256}.{ext}", "审核后的模块、转码和派生版本。"),
        ("brand-assets/preview/{assetId}/{version}/", "缩略图、低码率代理视频。"),
        ("brand-assets/analysis/{assetId}/{analysisVersion}/", "关键帧和分析产物。"),
        ("system-deploy/", "部署包；不参与素材统计。"),
    ], [6000, 3360])
    add_heading(doc, "6.1 不变原则", 2)
    for item in ["Bucket 固定：saidian-brand-assets-prod-sz。", "Region 固定：oss-cn-shenzhen。", "现有对象地址保持兼容，不移动、不删除、不重新上传。", "原始目录和OSS原件只读；派生内容写入独立目录。", "精确重复命中同一SHA256时，不创建第二项逻辑素材。"]:
        list_item(doc, item)

    add_page_break(doc)
    add_heading(doc, "七、上传与自动处理流水线", 1)
    add_heading(doc, "7.1 员工上传", 2)
    for item in [
        "Web端和手机浏览器每批最多20个文件，单文件上限200MB。",
        "API使用磁盘流式暂存，不把整个文件长期驻留Node.js内存；批次完成后删除临时文件。",
        "员工只填写来源、产品范围、素材类型、内容说明、原创状态、权限、获得日期和员工归属。",
        "产品型号从 Product 表选择；如果无法确认，保存 UNKNOWN/PENDING，不自动设为可用。",
        "大体积原始视频继续由本地只读扫描代理处理。",
    ]:
        list_item(doc, item, ordered=True)

    add_heading(doc, "7.2 处理顺序", 2)
    add_table(doc, ["步骤", "处理", "结果"], [
        ("1", "接收文件和员工字段", "UploadBatch / UploadEvent"),
        ("2", "计算SHA256", "精确重复复用主素材；保留本次上传行为"),
        ("3", "上传深圳OSS", "原件对象键、版本ID、ETag"),
        ("4", "技术参数", "Sharp / FFprobe：尺寸、时长、编码、质量基础分"),
        ("5", "预览", "缩略图、低码率代理视频"),
        ("6", "内容提取", "OCR、语音转写、关键帧、场景切分"),
        ("7", "内容理解", "产品、人物、场景、功能、痛点和模块建议"),
        ("8", "近似去重", "候选关系和相似度；不自动删除"),
        ("9", "人工审核", "通过、退回、仅内部或拒绝"),
        ("10", "进入素材池", "满足硬规则后 ACTIVE"),
    ], [900, 3300, 5160])

    add_heading(doc, "7.3 失败和重试", 2)
    add_body(doc, "AI任务失败时不删除素材。系统按1、5、30分钟重试；达到最大次数后进入FAILED和人工队列。百炼未配置时使用UNCONFIGURED，不计为本地工具失败。")

    add_page_break(doc)
    add_heading(doc, "八、AI能力设计", 1)
    add_table(doc, ["能力", "执行方", "配置状态规则"], [
        ("技术参数", "Sharp / FFprobe", "本地工具可用即 AVAILABLE"),
        ("缩略图/代理视频", "Sharp / FFmpeg", "本地工具和OSS可用"),
        ("关键帧/场景切分", "FFmpeg", "本地工具可用"),
        ("近似去重", "感知哈希", "先生成指纹，再计算候选"),
        ("图片理解/OCR", "阿里云百炼多模态", "BAILIAN_API_KEY + 模型配置后为 CONFIGURED；实际成功后记录结果"),
        ("语音转写", "阿里云百炼语音服务", "转写地址和模型均配置后为 CONFIGURED"),
        ("自动标签", "百炼建议 + 受控字典", "AI标签不覆盖 locked 人工标签"),
    ], [1900, 2700, 4760])

    add_heading(doc, "8.1 环境变量", 2)
    add_table(doc, ["变量", "用途"], [
        ("OSS_REGION / OSS_ENDPOINT", "固定深圳地域和端点"),
        ("OSS_BUCKET", "saidian-brand-assets-prod-sz"),
        ("OSS_ACCESS_KEY_ID / OSS_ACCESS_KEY_SECRET", "OSS最小权限RAM凭据"),
        ("BAILIAN_API_KEY", "百炼访问凭据；为空时显示未配置"),
        ("BAILIAN_BASE_URL", "OpenAI兼容接口地址"),
        ("BAILIAN_VISION_MODEL", "图片、关键帧、OCR和内容理解模型"),
        ("BAILIAN_TRANSCRIPTION_URL / MODEL", "语音转写接口和模型"),
    ], [3700, 5660])

    add_callout(doc, "能力显示规则", "“已配置”只表示参数完整；“可用/成功”必须有真实任务回执。接口未配置时不把Sharp、FFmpeg结果显示为AI识别完成。", "FFF4E5", RED)

    add_page_break(doc)
    add_heading(doc, "九、页面设计", 1)
    add_heading(doc, "9.1 品牌知识库", 2)
    add_table(doc, ["页面", "主要能力"], [
        ("知识条目", "按类型、型号、状态检索；编辑后重新审核；显示来源等级和AI调用资格。"),
        ("品牌版本", "品牌定位、故事、语气、视觉、国家版本和生效时间。"),
        ("产品库", "产品、型号、系列、SKU和证据关系。"),
        ("FAQ", "标准问题、不同问法、短回复、详细回复、频次和调用资格。"),
        ("证据/型号映射/表述规则", "保留现有底表并作为知识调用约束。"),
    ], [2100, 7260])

    add_heading(doc, "9.2 素材库", 2)
    add_table(doc, ["工作区", "主要能力"], [
        ("素材总览", "总量、今日新增、待审核、AI失败、高质量、可调用和缺口。"),
        ("批量上传", "20文件队列、员工归属、产品选择、重复结果、处理状态。"),
        ("素材检索", "型号、类型、层级、模块、员工、权限、评分和三类状态组合筛选。"),
        ("审核工作台", "预览、AI建议、人工字段、权限、通过/退回/仅内部/拒绝。"),
        ("视频模块", "时间轴、转写、切段建议、模块分类、锁定和高质量物化。"),
        ("素材详情", "版本、标签、来源、关系、审核、任务、使用和发布效果。"),
        ("AI任务", "执行方、模型、状态、尝试次数、失败或未配置原因。"),
        ("缺口与日报", "员工增量、型号覆盖、补拍建议、重复、AI派生和实际调用。"),
    ], [2100, 7260])

    add_page_break(doc)
    add_heading(doc, "十、API 合同", 1)
    add_table(doc, ["方法", "路径", "用途"], [
        ("POST", "/brand-data/upload-batches", "建立批次和员工/产品归属"),
        ("POST", "/brand-data/upload-batches/:id/files", "流式批量上传，最多20个文件"),
        ("GET", "/brand-data/upload-batches/:id", "批次进度、重复和失败明细"),
        ("GET", "/brand-data/assets", "游标分页和组合筛选；无分页参数保留V1数组兼容"),
        ("GET", "/brand-data/assets/:id", "完整素材对象图谱"),
        ("PATCH", "/brand-data/assets/:id/metadata", "结构化元数据、产品和权限"),
        ("POST", "/brand-data/assets/:id/review", "通过、退回、仅内部或拒绝"),
        ("POST", "/brand-data/assets/:id/reanalyze", "新分析版本，不覆盖人工锁定标签"),
        ("GET", "/brand-data/assets/:id/segments", "视频切段清单"),
        ("PATCH", "/brand-data/assets/:id/segments/:segmentId", "确认时间、转写和模块分类"),
        ("POST", "/brand-data/assets/:id/segments/:segmentId/materialize", "生成高质量模块文件和派生关系"),
        ("GET", "/brand-data/analysis-jobs", "AI处理队列"),
        ("GET", "/brand-data/asset-gaps", "素材缺口"),
        ("GET", "/brand-data/reports/daily", "每日素材经营台账"),
        ("GET", "/brand-data/ai-capabilities", "本地与百炼能力状态"),
    ], [900, 4800, 3660])

    add_heading(doc, "10.1 统一追溯字段", 2)
    add_body(doc, "所有素材响应统一提供：发生时间、素材编号、员工编号、执行主体、来源、处理状态、审核状态、可用状态、OSS对象、失败原因和审计编号。员工归属使用 Employee.id，姓名只作为展示。")

    add_page_break(doc)
    add_heading(doc, "十一、增量迁移与兼容", 1)
    add_body(doc, "现项目没有可执行的Prisma迁移基线，正式环境沿用已验证的 prisma db push 增量同步。V2增加可重复执行的 backfill-asset-v2.ts，用于结构化回填和对账。")
    add_table(doc, ["步骤", "动作", "保护措施"], [
        ("1", "备份试运行数据库并记录素材、版本、OSS对象数量", "只读对账，不改OSS原件"),
        ("2", "执行 Prisma schema 增量同步", "新字段使用默认值或可空字段"),
        ("3", "执行 V2 回填脚本", "可重复执行；公开编号、状态、权限、产品和标签增量补齐"),
        ("4", "从 model 和 sourceSnapshot 回填关系", "无法确认则 UNKNOWN/PENDING"),
        ("5", "FAQ从现有知识条目生成标准记录", "保留原KnowledgeEntry"),
        ("6", "运行数量、哈希和OSS对象对账", "异常停止切换，不移动对象"),
        ("7", "切换V2页面；保留V1接口一个版本周期", "旧上传、更新、审核和无分页查询继续可用"),
    ], [900, 4300, 4160])

    add_heading(doc, "11.1 回填映射", 2)
    add_table(doc, ["V1字段", "V2目标", "无法确认时"], [
        ("sourceSnapshot.assetNo", "Asset.assetNo", "按V2规则生成"),
        ("mediaType", "Asset.kind", "DOCUMENT"),
        ("status", "reviewStatus + availabilityStatus", "PENDING + INACTIVE"),
        ("copyrightStatus", "rightsStatus", "AUTH_REQUIRED"),
        ("model", "AssetProduct", "保留旧model，scope=UNKNOWN"),
        ("aiTags/scenarios/sellingPoints", "TagDefinition + AssetTag", "source=IMPORT，未确认"),
        ("AssetVersion旧字段", "文件名、扩展名、尺寸、时长", "从Asset补齐"),
    ], [3000, 3860, 2500])

    add_page_break(doc)
    add_heading(doc, "十二、每日台账与素材缺口", 1)
    add_heading(doc, "12.1 每日台账", 2)
    add_table(doc, ["区块", "必须包含"], [
        ("素材新增", "文件、素材编号、来源、上传员工、批次、OSS对象、审核状态。"),
        ("重复与失败", "精确重复次数、主素材、失败原因、重试和人工处理状态。"),
        ("AI派生", "切片、预览、关键帧、模块素材及其原片关系。"),
        ("员工贡献", "上传数、正式新增数、重复数、失败数、审核通过和被调用数。"),
        ("使用与发布", "被哪个脚本/成片/AI任务调用，平台、账号、内容编号和员工。"),
        ("效果", "播放、点赞、评论、收藏、咨询、订单、成交；未获取字段明确标注。"),
        ("缺口", "型号覆盖、缺少类型、补拍建议、负责人和完成状态。"),
    ], [2100, 7260])

    add_heading(doc, "12.2 统计口径", 2)
    for item in [
        "员工上传：UploadEvent总数。",
        "正式新增：UploadEvent.result=CREATED；精确重复不计新增。",
        "AI派生模块：来源为AI_DERIVED且形成独立Asset。",
        "审核通过：AssetReviewDecision.action=APPROVE。",
        "实际调用：AssetUsage记录数；不能用生成数量代替。",
        "效果数据：AssetMetricSnapshot；字段未取得时写入 unavailableFields，不按0计算。",
    ]:
        list_item(doc, item)

    add_page_break(doc)
    add_heading(doc, "十三、四周实施顺序", 1)
    add_table(doc, ["阶段", "交付", "完成定义"], [
        ("第1周\n模型和标准", "V2文档、枚举、关系表、Prisma schema、回填脚本、受控标签", "schema有效；旧数据可回填；无确认数据不自动ACTIVE"),
        ("第2周\n上传与管理", "批量上传、游标检索、详情、审核、员工贡献、兼容入口", "20文件批次可追溯；精确重复不产生第二主素材"),
        ("第3周\nAI处理", "技术参数、预览、OCR、转写、关键帧、切段、标签、近似去重", "已配置能力有任务回执；未配置项清楚显示"),
        ("第4周\n台账与缺口", "使用台账、员工日报、缺口分析、补拍建议、试运行切换", "数量和OSS对象对账；日报区分各统计口径"),
    ], [1800, 3900, 3660])

    add_heading(doc, "13.1 上线门槛", 2)
    for item in ["试运行数据库完成 schema 增量同步和V2回填。", "素材总数、版本数、OSS对象数完成对账。", "API、页面、构建和基础上传流程通过。", "至少一张图片完成上传、精确去重、技术分析和人工审核演练。", "百炼未配置时保持未配置；不得为了验收伪造成功状态。"]:
        list_item(doc, item, ordered=True)

    add_page_break(doc)
    add_heading(doc, "十四、测试与验收矩阵", 1)
    add_table(doc, ["编号", "场景", "预期结果"], [
        ("A01", "原始目录和OSS原件", "迁移前后SHA256一致，无移动和覆盖。"),
        ("A02", "同一文件不同名称上传", "一个主素材、两次UploadEvent，重复不计员工新增。"),
        ("A03", "近似图片/视频", "只生成NEAR_DUPLICATE候选，不自动删除或合并。"),
        ("A04", "多型号素材", "AssetProduct可关联多个产品并正常筛选。"),
        ("A05", "AI失败", "保留素材、任务、错误、尝试次数和下一次重试。"),
        ("A06", "人工修改AI标签", "AssetTag.locked=true，后续分析不得覆盖。"),
        ("A07", "视频片段", "可追溯原片、时间范围、分析版本和创建员工。"),
        ("A08", "内容调用资格", "只返回APPROVED + ACTIVE + COMMERCIAL/EDIT_ONLY。"),
        ("A09", "素材详情", "上传人、审核人、版本、OSS对象、关系和审计完整。"),
        ("A10", "AI处理覆盖", "支持格式90%以上完成或给出明确失败/未配置原因。"),
        ("A11", "图片时效", "20MB以内图片在正常负载下5分钟内完成分析。"),
        ("A12", "视频时效", "10分钟以内视频在正常负载下30分钟内完成转写和切段。"),
        ("A13", "日报口径", "员工上传、AI派生、重复、审核和实际调用分别统计。"),
        ("A14", "迁移对账", "素材、版本和OSS对象数量对账。"),
        ("A15", "工程质量", "Prisma校验、TypeScript、单测、API集成、生产构建和页面验证通过。"),
    ], [800, 3500, 5060])

    add_page_break(doc)
    add_heading(doc, "十五、运行与维护", 1)
    add_heading(doc, "15.1 每小时", 2)
    for item in ["处理待执行和到期重试的AssetAnalysisJob。", "刷新处理状态和失败原因。", "生成待审核任务；不自动改变人工锁定标签。"]:
        list_item(doc, item)
    add_heading(doc, "15.2 每天", 2)
    for item in ["统计员工上传、正式新增、重复、失败、审核和实际调用。", "计算型号与素材类型覆盖缺口。", "检查授权待确认、已过期、长时间未审核和AI失败项。", "汇总发布效果；未获取值保持未获取。"]:
        list_item(doc, item)
    add_heading(doc, "15.3 运维检查", 2)
    add_table(doc, ["检查项", "正常标准"], [
        ("PostgreSQL", "连接正常、任务可写入、无长时间锁。"),
        ("深圳OSS", "Bucket健康、原件可签名读取、派生目录可写。"),
        ("FFmpeg/FFprobe", "可读取视频参数并生成预览/切段。"),
        ("Sharp", "可读取图片参数、生成缩略图和感知哈希。"),
        ("百炼", "配置、模型、真实任务回执和最近成功时间分开显示。"),
        ("临时目录", "批次结束后文件被清理；失败仍保留数据库记录。"),
    ], [2800, 6560])

    add_page_break(doc)
    add_heading(doc, "十六、交付清单", 1)
    add_table(doc, ["交付物", "位置/说明"], [
        ("V2设计文档", str(OUTPUT)),
        ("V1原规范", f"保持只读：{SOURCE}"),
        ("数据模型", str(REPO / "apps/api/prisma/schema.prisma")),
        ("V2回填", str(REPO / "apps/api/prisma/backfill-asset-v2.ts")),
        ("AI处理服务", str(REPO / "apps/api/src/asset-ai.service.ts")),
        ("品牌数据API", str(REPO / "apps/api/src/brand-data.service.ts")),
        ("管理页面", str(REPO / "apps/admin/src/components/BrandDataCenter.vue")),
        ("配置模板", str(REPO / ".env.example")),
    ], [2600, 6760])
    add_callout(doc, "最终原则", "素材库不是“存文件的地方”，而是可检索、可组合、可追溯、可复盘、可被AI安全调用的品牌内容资产中心。", "E8EEF5", NAVY)

    doc.core_properties.title = "Saydian AI 素材库 V2.0 落地设计与实施方案"
    doc.core_properties.subject = "赛电品牌数据中心：品牌知识库与素材库 V2.0"
    doc.core_properties.author = "Saydian AI OS"
    doc.core_properties.keywords = "Saydian, 素材库, 品牌知识库, OSS, AI处理, Prisma"
    doc.core_properties.comments = "V1文档保持只读，本文件为独立V2实施方案。"
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
